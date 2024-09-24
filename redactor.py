import spacy
import re
from docx import Document
import uuid

# Load the fine-tuned spaCy model (Replace with your custom model path)
try:
    nlp = spacy.load("path_to_your_fine_tuned_model")  # Fine-tuned model
except:
    nlp = spacy.load("en_core_web_trf")  # Fallback to default model

# Global token map for consistent replacement
token_map = {}

def generate_token():
    return f"[[REDACTED_{uuid.uuid4().hex[:8]}]]"

def generate_token_for_term(term):
    if term not in token_map:
        token_map[term] = generate_token()
    return token_map[term]

def analyze_document(filepath):
    doc = Document(filepath)
    text = "\n".join([para.text for para in doc.paragraphs])
    return suggest_redactions(text)

# Add more custom patterns if needed
def suggest_redactions(text):
    suggestions = {}

    # Process the text with the NLP model
    doc = nlp(text)

    for ent in doc.ents:
        # Assume the model provides confidence scores; you might need to adjust this based on your model
        confidence = getattr(ent._, 'confidence', 0.9)  # Placeholder confidence score
        if confidence >= 0.7:  # Use a confidence threshold
            if ent.label_ in ['PERSON', 'ORG', 'GPE', 'DATE', 'MONEY', 'PERCENT', 'QUANTITY', 'PROJECT_NAME', 'COMPANY', 'POSITION_TITLE']:
                suggestions[ent.text] = {
                    'type': ent.label_,
                    'context': get_context(text, ent.start_char, ent.end_char),
                    'confidence': confidence,
                    'token': generate_token_for_term(ent.text)
                }

    # Add custom patterns for sensitive info like SSNs, Credit Cards, etc.
    patterns = [
        (r'\b\d{3}-\d{2}-\d{4}\b', 'SSN'),  # Social Security Number
        (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 'EMAIL'),
        # Add more patterns as needed
    ]

    for pattern, label in patterns:
        for match in re.finditer(pattern, text):
            suggestions[match.group()] = {
                'type': label,
                'context': get_context(text, match.start(), match.end()),
                'confidence': 1.0,  # Regular expressions are certain matches
                'token': generate_token_for_term(match.group())
            }

    return suggestions

def get_context(text, start, end, context_size=50):
    context_start = max(0, start - context_size)
    context_end = min(len(text), end + context_size)
    return text[context_start:context_end] + '...'

def redact_document(filepath, approved_redactions, redaction_id):
    doc = Document(filepath)
    redaction_map = {}

    for para in doc.paragraphs:
        for original, info in approved_redactions.items():
            if original in para.text and info['action'] != 'IGNORE':
                if info['action'] == 'CUSTOM':
                    replacement = info.get('custom_value', generate_token_for_term(original))
                elif info['action'] == 'MASK':
                    replacement = 'X' * len(original)
                else:  # REDACT
                    replacement = generate_token_for_term(original)
                para.text = para.text.replace(original, replacement)
                redaction_map[replacement] = {'original': original, 'type': info['type']}

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for original, info in approved_redactions.items():
                    if original in cell.text and info['action'] != 'IGNORE':
                        if info['action'] == 'CUSTOM':
                            replacement = info.get('custom_value', generate_token_for_term(original))
                        elif info['action'] == 'MASK':
                            replacement = 'X' * len(original)
                        else:  # REDACT
                            replacement = generate_token_for_term(original)
                        cell.text = cell.text.replace(original, replacement)
                        redaction_map[replacement] = {'original': original, 'type': info['type']}

    # Process headers and footers
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs + footer.paragraphs:
            for original, info in approved_redactions.items():
                if original in paragraph.text and info['action'] != 'IGNORE':
                    if info['action'] == 'CUSTOM':
                        replacement = info.get('custom_value', generate_token_for_term(original))
                    elif info['action'] == 'MASK':
                        replacement = 'X' * len(original)
                    else:  # REDACT
                        replacement = generate_token_for_term(original)
                    paragraph.text = paragraph.text.replace(original, replacement)
                    redaction_map[replacement] = {'original': original, 'type': info['type']}

    redacted_filepath = filepath.replace('.docx', f'_{redaction_id}_redacted.docx')
    doc.save(redacted_filepath)
    return redacted_filepath, redaction_map

def restore_document(redacted_filepath, redaction_map):
    doc = Document(redacted_filepath)

    for para in doc.paragraphs:
        for token, info in redaction_map.items():
            para.text = para.text.replace(token, info['original'])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for token, info in redaction_map.items():
                    cell.text = cell.text.replace(token, info['original'])

    # Restore headers and footers
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs + footer.paragraphs:
            for token, info in redaction_map.items():
                paragraph.text = paragraph.text.replace(token, info['original'])

    restored_filepath = redacted_filepath.replace('_redacted.docx', '_restored.docx')
    doc.save(restored_filepath)
    return restored_filepath

def get_preview(filepath, approved_redactions):
    doc = Document(filepath)
    preview_text = []

    for para in doc.paragraphs:
        redacted_para = para.text
        for original, info in approved_redactions.items():
            if original in redacted_para and info['action'] != 'IGNORE':
                if info['action'] == 'CUSTOM':
                    replacement = info.get('custom_value', generate_token_for_term(original))
                elif info['action'] == 'MASK':
                    replacement = 'X' * len(original)
                else:  # REDACT
                    replacement = generate_token_for_term(original)
                redacted_para = redacted_para.replace(original, replacement)
        preview_text.append(redacted_para)

    return "\n".join(preview_text)

def search_document(filepath, query):
    doc = Document(filepath)
    results = []
    for i, para in enumerate(doc.paragraphs):
        if query.lower() in para.text.lower():
            start = max(0, para.text.lower().index(query.lower()) - 20)
            end = min(len(para.text), para.text.lower().index(query.lower()) + len(query) + 20)
            context = para.text[start:end]
            results.append({
                'paragraph': i + 1,
                'context': context
            })
    return results
